"""
Extracts raw text from resume input: PDF, DOCX, or plain pasted text.

Falls back to OCR (Tesseract) when a PDF page has little to no extractable
text layer, which usually means it's a scanned image rather than a text PDF.

Note: the OCR fallback depends on TWO separate system binaries (not pip
packages):
  - poppler (pdf2image shells out to it, to rasterize PDF pages to images)
  - Tesseract (pytesseract shells out to it, to read text from those images)
Either one can be missing independently. On Linux: `apt-get install
poppler-utils tesseract-ocr`. On Windows: download a poppler build (e.g.
github.com/oschwartz10612/poppler-windows/releases) and a Tesseract
installer (e.g. github.com/UB-Mannheim/tesseract/wiki), then add both
bin/ folders to PATH and restart your terminal/IDE (PATH changes don't
apply to already-running processes). If either is missing, this raises
OcrUnavailableError with that guidance rather than letting the raw
pdf2image/pytesseract exception leak through unexplained.
"""

import io

import fitz  # PyMuPDF
import pytesseract
from docx import Document
from pdf2image import convert_from_bytes
from pdf2image.exceptions import PDFInfoNotInstalledError
from PIL import Image
from pytesseract import TesseractNotFoundError

# If a page yields fewer than this many characters via direct text extraction,
# treat it as likely scanned and fall back to OCR for that page.
MIN_CHARS_PER_PAGE = 20


class UnsupportedFileTypeError(Exception):
    """Raised when the uploaded file is neither a PDF nor a DOCX."""


class OcrUnavailableError(Exception):
    """Raised when OCR fallback is needed but poppler and/or Tesseract isn't installed/on PATH."""


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text from a PDF, page by page, falling back to OCR for pages
    that appear to be scanned images rather than text.
    """
    pages_text: list[str] = []
    doc = fitz.open(stream=file_bytes, filetype="pdf")

    pages_needing_ocr: list[int] = []
    for page_index, page in enumerate(doc):
        text = page.get_text().strip()
        if len(text) < MIN_CHARS_PER_PAGE:
            pages_needing_ocr.append(page_index)
            pages_text.append("")  # placeholder, filled in below
        else:
            pages_text.append(text)
    doc.close()

    if pages_needing_ocr:
        try:
            ocr_images = convert_from_bytes(file_bytes)
        except PDFInfoNotInstalledError as exc:
            raise OcrUnavailableError(
                "This PDF has page(s) with little to no extractable text and "
                "needs OCR, but poppler (required by pdf2image) isn't "
                "installed or isn't on PATH. On Linux: `apt-get install "
                "poppler-utils`. On Windows: download a poppler build and "
                "add its bin/ folder to PATH, then restart your terminal/IDE."
            ) from exc

        for page_index in pages_needing_ocr:
            if page_index < len(ocr_images):
                try:
                    pages_text[page_index] = _ocr_image(ocr_images[page_index])
                except TesseractNotFoundError as exc:
                    raise OcrUnavailableError(
                        "This PDF has page(s) with little to no extractable text "
                        "and needs OCR, but Tesseract isn't installed or isn't on "
                        "PATH. On Linux: `apt-get install tesseract-ocr`. On "
                        "Windows: install from github.com/UB-Mannheim/tesseract/wiki "
                        "and add its install folder to PATH, then restart your "
                        "terminal/IDE."
                    ) from exc

    return "\n\n".join(p for p in pages_text if p)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a .docx file, including paragraphs and table cells."""
    document = Document(io.BytesIO(file_bytes))

    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    return "\n".join(parts)


def _ocr_image(image: Image.Image) -> str:
    """Run Tesseract OCR on a single page image."""
    return pytesseract.image_to_string(image).strip()


def extract_text(
    *,
    file_bytes: bytes | None,
    filename: str | None,
    raw_text: str | None,
) -> str:
    """
    Single entry point for extraction. Exactly one of (file_bytes + filename)
    or raw_text should be provided; the caller (API layer) is responsible
    for that validation.
    """
    if raw_text is not None:
        return raw_text.strip()

    if file_bytes is None or filename is None:
        raise ValueError("Either raw_text or (file_bytes and filename) must be provided.")

    lower_name = filename.lower()
    if lower_name.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif lower_name.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    else:
        raise UnsupportedFileTypeError(
            f"Unsupported file type for '{filename}'. Only .pdf and .docx are supported."
        )